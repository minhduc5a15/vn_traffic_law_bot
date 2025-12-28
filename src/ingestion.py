import os
import re
from docx import Document
from langchain_core.documents import Document as LangchainDocument
from typing import List, Dict


class VietnameseLawParser:
    def __init__(self, data_path: str = "data/raw"):
        self.data_path = data_path

    def load_and_parse(self) -> List[LangchainDocument]:
        all_documents = []
        if not os.path.exists(self.data_path):
            return []

        files = [f for f in os.listdir(self.data_path) if f.endswith(".docx")]
        print(
            f"🔄 [V4-Ultimate] Đang xử lý {len(files)} file (Fix lỗi Title & Context)..."
        )

        for file_name in files:
            file_path = os.path.join(self.data_path, file_name)
            try:
                docs = self._process_single_file(file_path, file_name)
                all_documents.extend(docs)
                print(f"   -> ✅ {file_name}: {len(docs)} chunks")
            except Exception as e:
                print(f"   -> ❌ Lỗi {file_name}: {e}")

        return all_documents

    def _get_doc_type_and_name(self, file_name: str, doc_content: List[str]) -> Dict:
        name_lower = file_name.lower()
        doc_type = "van_ban_khac"
        law_id = file_name.replace(".docx", "")

        if "luat" in name_lower or re.match(r"^\d+_\d+_qh", name_lower):
            doc_type = "luat"
        elif "nd-cp" in name_lower or "nghi_dinh" in name_lower:
            doc_type = "nghi_dinh"
        elif "tt-" in name_lower or "thong_tu" in name_lower:
            doc_type = "thong_tu"

        law_name = law_id
        # Cố gắng bắt chính xác số hiệu văn bản từ nội dung
        for line in doc_content[:30]:  # Quét sâu hơn chút
            line = line.strip()
            # Bắt dạng: "Số: 168/2024/NĐ-CP" hoặc "Luật số 36..."
            if re.search(r"(Số|Luật số)[:\s]+(\d+[\/\-].*)", line, re.IGNORECASE):
                law_name = line.replace("Số:", "").replace("Luật số:", "").strip()
                break

        if law_name == law_id:
            # VD: 168_2024_ND-CP -> Nghị định 168/2024/ND-CP
            if "ND-CP" in law_id:
                law_name = f"Nghị định {law_id.split('_')[0]}/2024/NĐ-CP"

        return {"doc_type": doc_type, "law_name": law_name, "law_id": law_id}

    def _process_single_file(
        self, file_path: str, file_name: str
    ) -> List[LangchainDocument]:
        doc = Document(file_path)
        # Bỏ dòng trống ngay từ đầu
        full_lines = [p.text.strip() for p in doc.paragraphs if p.text.strip()]
        doc_info = self._get_doc_type_and_name(file_name, full_lines)

        documents = []

        current_article = None
        current_clause = None

        article_re = re.compile(r"^Điều\s+(\d+)[\.:]?\s*(.*)$", re.IGNORECASE)
        clause_re = re.compile(r"^(\d+)\.\s+(.*)$")
        point_re = re.compile(r"^([a-zđ])\)\s+(.*)$", re.IGNORECASE)

        def commit_clause():
            if not current_clause or not current_article:
                return

            clause_intro = "\n".join(current_clause["content_lines"])

            # Thêm Tên Luật vào đầu mỗi chunk
            context_header = f"{doc_info['law_name']} > {current_article['full_title']}"

            # 1. Parent Chunk
            full_text = (
                f"{context_header}\nKhoản {current_clause['id']}: {clause_intro}"
            )
            for p in current_clause["points"]:
                full_text += f"\nĐiểm {p['id']}) {p['content']}"

            documents.append(
                LangchainDocument(
                    page_content=full_text,
                    metadata={
                        **doc_info,
                        "article": current_article["id"],
                        "clause": current_clause["id"],
                        "point": "all",
                        "citation": f"{doc_info['law_name']} - Điều {current_article['id']} Khoản {current_clause['id']}",
                        "is_parent": True,
                    },
                )
            )

            # 2. Child Chunks
            for p in current_clause["points"]:
                # Inject full context: Luật > Điều > Khoản > Điểm
                enriched_content = (
                    f"{context_header}\n"
                    f"Khoản {current_clause['id']}: {clause_intro}\n"
                    f"Điểm {p['id']}) {p['content']}"
                )

                documents.append(
                    LangchainDocument(
                        page_content=enriched_content,
                        metadata={
                            **doc_info,
                            "article": current_article["id"],
                            "clause": current_clause["id"],
                            "point": p["id"],
                            "citation": f"{doc_info['law_name']} - Điều {current_article['id']} Khoản {current_clause['id']} Điểm {p['id']}",
                            "is_parent": False,
                        },
                    )
                )

        i = 0
        while i < len(full_lines):
            line = full_lines[i]

            # 1. Phát hiện ĐIỀU
            art_match = article_re.match(line)
            if art_match:
                commit_clause()
                current_clause = None

                art_id = art_match.group(1)
                title_part = art_match.group(2).strip()

                # Nếu tiêu đề ngắn (rỗng hoặc chỉ có dấu chấm), hoặc dòng tiếp theo viết hoa toàn bộ -> Nối dòng tiếp theo vào
                full_title = f"Điều {art_id}. {title_part}"

                # Kiểm tra dòng tiếp theo (Lookahead)
                if i + 1 < len(full_lines):
                    next_line = full_lines[i + 1]
                    # Logic heuristic: Nếu dòng sau không phải Khoản (1.), không phải Điểm (a)), và không phải Điều mới
                    # Thường tiêu đề Điều viết hoa hoặc đậm, nhưng ở đây ta check đơn giản:
                    if (
                        not clause_re.match(next_line)
                        and not article_re.match(next_line)
                        and not point_re.match(next_line)
                    ):
                        # Giả định đó là phần tiếp theo của tiêu đề
                        full_title += " " + next_line
                        i += 1  # Bỏ qua dòng tiếp theo vì đã gộp rồi

                current_article = {"id": art_id, "full_title": full_title}
                current_clause = {"id": "intro", "content_lines": [], "points": []}
                i += 1
                continue

            if not current_article:
                i += 1
                continue

            # 2. Phát hiện KHOẢN
            clause_match = clause_re.match(line)
            if clause_match:
                commit_clause()
                current_clause = {
                    "id": clause_match.group(1),
                    "content_lines": [clause_match.group(2).strip()],
                    "points": [],
                }
                i += 1
                continue

            # 3. Phát hiện ĐIỂM
            point_match = point_re.match(line)
            if point_match and current_clause:
                current_clause["points"].append(
                    {
                        "id": point_match.group(1).lower(),
                        "content": point_match.group(2).strip(),
                    }
                )
                i += 1
                continue

            # 4. Nội dung nối tiếp
            if current_clause:
                if current_clause["points"]:
                    current_clause["points"][-1]["content"] += " " + line
                else:
                    current_clause["content_lines"].append(line)
            i += 1

        commit_clause()
        return documents
